import email, smtplib, ssl 
from email import encoders 
from email.mime.base import MIMEBase 
from email.mime.multipart import MIMEMultipart 
from email.mime.text import MIMEText
from docx import Document
from docx.shared import Cm, Inches
from docx.shared import Pt
from datetime import date, datetime
import datetime
import sys
import pytz
import senderParticulars 
import Blk9_2Rm_BuildSani
import Blk9_2Rm_Elec
import Blk165a_2Rm_BuildSani
import Blk165a_2Rm_Elec
import Blk165a_1Rm_BuildSani
import Blk165a_1Rm_Elec
import Blk536_2Rm_BuildSani
import Blk536_2Rm_Elec
import Blk536_1Rm_BuildSani 
import Blk536_1Rm_Elec
import Blk632a_2Rm_BuildSani
import Blk632a_2Rm_Elec
import Blk632a_1Rm_BuildSani
import Blk632a_1Rm_Elec

# send email with attached works order to team C EE & notice of account closure to HS 
def sendemail(fn, receiverEmail, receiver_email_HS, blk, unitNo, typeOfWorks, streetName):
    sender_email = senderParticulars.email
    #team C AE email
    receiver_email = receiverEmail
    #HS team AE email
    receiverHS = receiver_email_HS
    # Building or Electrical WO
    type_Works = typeOfWorks
    # string to store the body of the email to CR AE
    body = senderParticulars.emailContent_CR
    #sender email password stored in a separate file
    password = senderParticulars.pw
    
    # instance of MIMEMultipart - to Team C AE
    msg = MIMEMultipart()  
    msg['From'] = sender_email
    # storing the receivers email address - CR AE to issue WO  
    msg['To'] = receiver_email  
    msg['Subject'] = senderParticulars.subject_CR
    msg["Cc"] = senderParticulars.workEmail
    # attach the body with the msg instance - for issuance of WO 
    msg.attach(MIMEText(body, 'plain'))
    # open the file to be sent  
    filename = fn 
    # #Open the file in binary mode
    with open(filename, "rb") as attachment:
        #Add file as application/octet-stream
        #Email client can usually download this automatically as attachment
        part = MIMEBase("application", "octet-stream")
        part.set_payload(attachment.read())
    # #encode file in ASCII characters to send by email
    encoders.encode_base64(part)
    # # Add header as key/value pair to attachment part
    part.add_header(
        "Content-Disposition",
        f"attachment; filename= {filename}",
    )
    # Add attachment to message
    msg.attach(part)

    # instance of MIMEMultipart - to HS Finance
    msg_to_HS = MIMEMultipart('alternative')
    msg_to_HS['From'] = sender_email
    # storing the receivers email address - HS AE to close account 
    msg_to_HS['To'] = receiver_email_HS 
    msg_to_HS['Subject'] = senderParticulars.subject_HS
    # Cc CR EEIC, HS EE & HS EM
    msg_to_HS["Cc"] = senderParticulars.workEmail + ";" + senderParticulars.hs_Rental_Ee + ";" + senderParticulars.hs_Rental_Em
    plaintext = senderParticulars.message_finance_plaintext.format(blk=blk, street=streetName, unitNo=unitNo)
    htmltext = senderParticulars.message_finance_html.format(blk=blk, street=streetName, unitNo=unitNo)
    msg_to_HS.attach(MIMEText(plaintext, 'plain'))
    msg_to_HS.attach(MIMEText(htmltext, 'html'))
    
    # Log in to server using secure context and send email
    context = ssl.create_default_context()
    with smtplib.SMTP_SSL("smtp.gmail.com", 465, context=context) as server:
        server.login(sender_email, password)
        # Send to Team C AE for issuing of WO
        server.sendmail(sender_email, (receiver_email,senderParticulars.workEmail),msg.as_string())
        if (type_Works == "B"):
            # Send to HS AE for closing of rental account (For Building WO only - business rule)
            server.sendmail(sender_email, (receiverHS,senderParticulars.workEmail, senderParticulars.hs_Rental_Em, senderParticulars.hs_Rental_Ee),\
                msg_to_HS.as_string())

#set Works order Commencement & Delivery date
def cAndD_Date(date):
	while True:
		#test if date is within monday - friday
		if(date.weekday() < 5):
			break
		else:
			date += datetime.timedelta(days=1)
	return date.strftime("%d/%m/%Y")

#set word document columns width
def set_col_widths(table):
    widths = (Inches(1), Inches(1.2), Inches(3.5))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

#sor item details
def sorItem(blkRoomTypeOfWorks, today, document):
    # Standard template headers for instruction orders -------------
    items = (
    	("", "Commencement Date", cAndD_Date(today + datetime.timedelta(days=12)), "Delivery Date", cAndD_Date(today + datetime.timedelta(days=25)), ""),
    	("", "Completion Date", cAndD_Date(today + datetime.timedelta(days=25)), "", "", ""),
    )

    # add table 
    table_header = document.add_table(rows=1, cols=6,style="Table Grid")

    #populate header row 
    heading_cells = table_header.rows[0].cells
    heading_cells[0].text = ""
    heading_cells[1].text = 'Flat type (To indicate Sold/Vacant/Rental)'
    heading_cells[2].text = 'Rental'
    heading_cells[3].text = 'Contract No'
    heading_cells[4].text = blkRoomTypeOfWorks.contractNo
    heading_cells[5].text = ""

    # add a data row for each item
    for i in items:
        cells = table_header.add_row().cells
        cells[0].text = i[0]
        cells[1].text = i[1]
        cells[2].text = i[2]
        cells[3].text = i[3]
        cells[4].text = i[4]
        cells[5].text = i[5]

    # add table ------------------
    items_table = document.add_table(rows=1, cols=6,style="Table Grid")

    # populate header row --------
    heading_cells = items_table.rows[0].cells
    heading_cells[0].text = "S/N"
    heading_cells[1].text = 'Description of works (eg. Location)'
    heading_cells[2].text = 'SOR'
    heading_cells[3].text = 'Quantity'
    heading_cells[4].text = 'Job code'
    heading_cells[5].text = 'Tax code'

    sn = 1
    # add a data row for each item
    for i in blkRoomTypeOfWorks.items:
        cells = items_table.add_row().cells
        cells[0].text = str(sn)
        cells[1].text = i[0]
        cells[2].text = i[1]
        cells[3].text = i[2]
        cells[4].text = i[3]
        cells[5].text = i[4]
        sn+=1

#generate works order with block, room, type of works - (brtype)
def generateWorksOrder(brtype, document):
    #staff name - refer to file that contain sender's email password
    eeName = senderParticulars.name

    p1 = document.add_paragraph()
    p1_word = "Instruction To Issue Purchase Order"
    runner_p1 = p1.add_run(p1_word)
    runner_p1.bold = True
    runner_p1.underline =True
    runner_p1.font.name = 'Arial'
    runner_p1.font.size = Pt(16)

    p2 = document.add_paragraph()
    p2_word = "Submitted by Name/Designation: "

    runner_p2 = p2.add_run(p2_word)
    runner_p2.bold = True
    runner_p2_1 = p2.add_run(eeName)
    runner_p2_1.bold = True
    runner_p2_1.underline = True
    fontP2 = runner_p2.font
    fontP2_1 = runner_p2_1.font
    fontP2.size = Pt(12)
    fontP2_1.size = Pt(12)

    p3 = document.add_paragraph()
    p3_word = "Signature: "
    runner_p3 = p3.add_run(p3_word)
    runner_p3.bold = True
    runner_p3_1 = p3.add_run(eeName)
    runner_p3_1.bold = True
    runner_p3_1.underline = True
    font_p3 = runner_p3.font
    font_p3.size = Pt(12)
    font_p3_1 = runner_p3_1.font
    font_p3_1.size = Pt(12)

    p4 = document.add_paragraph()
    sg = pytz.timezone("Asia/Singapore")
    today = datetime.datetime.now(sg)
    today_date = today.strftime("%d/%m/%Y")
    p4_word = "Date: " + str(today_date)
    runner_p4 = p4.add_run(p4_word)
    runner_p4.bold = True
    font_p4 = runner_p4.font
    font_p4.size = Pt(12)

    table = document.add_table(rows=2, cols=3, style="Table Grid")
    set_col_widths(table)

    Blk = table.cell(0,0).paragraphs[0].add_run("Blk")
    Blk.bold=True
    Blk.font.size = Pt(12)
    uN = table.cell(0,1).paragraphs[0].add_run("Unit No.")
    uN.bold=True
    uN.font.size = Pt(12)
    sN = table.cell(0,2).paragraphs[0].add_run("Street Name")
    sN.bold=True
    sN.font.size = Pt(12)
    table.cell(1,0).text = sys.argv[1]
    table.cell(1,1).text = sys.argv[2]
    table.cell(1,2).text = brtype.StreetName

    # WO header title
    p5 = document.add_paragraph()
    runner_p5 = p5.add_run(brtype.woTitle)
    runner_p5.font.size = Pt(14)

    # Works order SOR rate items
    sorItem(brtype, today, document)
    filename = sys.argv[1] + "_" + sys.argv[2]+ "_" + sys.argv[4] + ".docx"
    document.save(filename)
    d=dict()
    d['fileName'] = filename
    d['streetName'] = brtype.StreetName
    return d

#display confirmation message
def displayEmailSent(buildingOrSanitaryWorks,receiver_Email_WO, receiver_email_HS):
    if (buildingOrSanitaryWorks == "B"):
        print(f"\nEmail sent to {receiver_Email_WO} for Issuance of Work Order !\n\n\
Email sent to {receiver_email_HS} for Closing of Rental Account !")
    else: # For electrical WO, Don't have to inform HS (to prevent duplicate case)
        print(f"\nEmail sent to {receiver_Email_WO} for Issuance of Work Order !")

def main():
    try:
        if len(sys.argv) == 5 and sys.argv[3].isdigit():
            print("Starting the program...")
            document = Document()
            if(sys.argv[1] == '9' and sys.argv[3] == '2' and sys.argv[4] == "B"):
                file_street_name = generateWorksOrder(Blk9_2Rm_BuildSani, document)
            if(sys.argv[1] == '9' and sys.argv[3] == '2' and sys.argv[4] == "E"):          
                file_street_name = generateWorksOrder(Blk9_2Rm_Elec, document)    
            if(sys.argv[1].lower() == '165a' and sys.argv[3] == '2' and sys.argv[4] == "B"):
                file_street_name = generateWorksOrder(Blk165a_2Rm_BuildSani, document) 
            if(sys.argv[1].lower() == '165a' and sys.argv[3] == '2' and sys.argv[4] == "E"):
                file_street_name = generateWorksOrder(Blk165a_2Rm_Elec, document) 
            if(sys.argv[1].lower() == '165a' and sys.argv[3] == '1' and sys.argv[4] == "B"):
                file_street_name = generateWorksOrder(Blk165a_1Rm_BuildSani, document)  
            if(sys.argv[1].lower() == '165a' and sys.argv[3] == '1' and sys.argv[4] == "E"):
                file_street_name = generateWorksOrder(Blk165a_1Rm_Elec, document) 
            if((sys.argv[1] == '536' or sys.argv[1] == '537')  and sys.argv[3] == '2' and sys.argv[4] == "B"):
                file_street_name = generateWorksOrder(Blk536_2Rm_BuildSani, document)
            if((sys.argv[1] == '536' or sys.argv[1] == '537') and sys.argv[3] == '2' and sys.argv[4] == "E"):
                file_street_name = generateWorksOrder(Blk536_2Rm_Elec, document)   
            if((sys.argv[1] == '536' or sys.argv[1] == '537') and sys.argv[3] == '1' and sys.argv[4] == "B"):
                file_street_name = generateWorksOrder(Blk536_1Rm_BuildSani, document) 
            if((sys.argv[1] == '536' or sys.argv[1] == '537') and sys.argv[3] == '1' and sys.argv[4] == "E"):
                file_street_name = generateWorksOrder(Blk536_1Rm_Elec, document)  
            if((sys.argv[1].lower() == '632a' or sys.argv[1].lower() == '632b') and sys.argv[3] == '2' and sys.argv[4] == "B"):
                file_street_name = generateWorksOrder(Blk632a_2Rm_BuildSani, document)   
            if((sys.argv[1].lower() == '632a' or sys.argv[1].lower() == '632b') and sys.argv[3] == '2' and sys.argv[4] == "E"):
                file_street_name = generateWorksOrder(Blk632a_2Rm_Elec, document)  
            if((sys.argv[1].lower() == '632a' or sys.argv[1].lower() == '632b') and sys.argv[3] == '1' and sys.argv[4] == "B"):
                file_street_name = generateWorksOrder(Blk632a_1Rm_BuildSani, document)    
            if((sys.argv[1].lower() == '632a' or sys.argv[1].lower() == '632b') and sys.argv[3] == '1' and sys.argv[4] == "E"):
                file_street_name = generateWorksOrder(Blk632a_1Rm_Elec, document)
        f_n = file_street_name['fileName']
        s_n = file_street_name['streetName']
        
        if (sys.argv[4] == "B"):
            print(f"Generating WO for blk {sys.argv[1]}, {sys.argv[3]}-Room, Building works...")
        elif (sys.argv[4] == "E"):
            print(f"Generating WO for blk {sys.argv[1]}, {sys.argv[3]}-Room, Electrical works...")

        # sending to different AE's from CR & HS team for different estates    
        if (sys.argv[1].lower() == '632a' or sys.argv[1].lower() == '632b'):
            receiver_Email_WO = senderParticulars.teamC_bp_receiverEmail
            receiver_email_HS = senderParticulars.hS_bp_ae_email
        else:
            receiver_Email_WO = senderParticulars.teamC_ck_receiverEmail
            receiver_email_HS = senderParticulars.hS_ck_ae_email
        
        sendemail(f_n, receiver_Email_WO, receiver_email_HS, sys.argv[1], sys.argv[2], sys.argv[4], s_n)
        displayEmailSent(sys.argv[4], receiver_Email_WO, receiver_email_HS)

    except UnboundLocalError as error:
        print("Invalid inputs !\nPlease enter only rental blocks, 536/537/165a/9/632a/632b")
    #   block | unit no | 1 or 2 room flat | Building or Electrical works | 
        print("Please enter in this format: autoemail.py [blk] [unit no.] [1/2] [B/E]")
    except smtplib.SMTPException:
        print("Error sending email. SMTPException")
    except Exception as e:
        print("Error sending e-mail. Please contact system administrator !")

if __name__ == "__main__":
    main()